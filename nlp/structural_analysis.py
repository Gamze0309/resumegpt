import re
import spacy
import pdfplumber
import io
import fitz
from docx import Document

nlp = spacy.load("en_core_web_sm")

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"

def structural_analysis(parsed_sections, content, file_type):
    result = get_result_of_found_sections(parsed_sections)
    if file_type == PDF_MIME:
        result += get_structure_pdf(content)
    elif file_type == DOCX_MIME:
        result += get_structure_docx(content)
    return result

def get_result_of_found_sections(parsed_sections):
    mandatory_sections = ["experience", "education", "skills"]
    recommended_sections = ["projects", "certifications", "languages"]
    missing_sections = []
    missing_recommended_sections = []
    result = []

    section_names = parsed_sections.keys()

    # missing_sections includes missing mandatory sections.
    for required in mandatory_sections:
        if not any(required in section_name for section_name in section_names):
            missing_sections.append(required)

    if len(missing_sections) > 0:
        result.append("Mandatory sections are missing: " + ", ".join(missing_sections) + ".")
    else:
        result.append("You provided mandatory sections.")

    # missing_recommended_sections includes missing recommended sections.
    for recommended in recommended_sections:
        if not any(recommended in section_name for section_name in section_names):
            missing_recommended_sections.append(recommended)

    if len(missing_recommended_sections) > 0:
        result.append("Some recommended sections are missing: " + ", ".join(missing_recommended_sections) + ".")
    else:
        result.append("You provided recommended sections (projects, certifications and languages).")
    
    # Check for the all sections are not blank
    for section, content in parsed_sections.items():
        if not bool(content.strip()):
            result.append(section + " section content are missing.")

    return result

def get_structure_pdf(pdf_content):
    result = []
    with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
        for page in pdf.pages:
            words = page.extract_words()
            x_positions = [float(word['x0']) for word in words]

            if not x_positions:
                continue

            x_positions.sort()
            threshold = 50
            clusters = [[x_positions[0]]]

            for x in x_positions[1:]:
                if abs(x - clusters[-1][-1]) > threshold:
                    clusters.append([x])
                else:
                    clusters[-1].append(x)
            if len(clusters) >= 2:
                result.append("The resume structure is multi-column. ATS may not be able to access the content in the multi-column structure.")

    doc = fitz.open(stream=io.BytesIO(pdf_content), filetype="pdf")
    if has_vector_graphics_pdf(doc):
        result.append("The resume contains some graphic or drawings. ATS may not be able to access the content in the graphics.")

    return result

def has_vector_graphics_pdf(doc) -> bool: 
    for page in doc: 
        drawings = page.get_drawings()
        if len(drawings) > 0: 
            return True 
    return False

def get_structure_docx(docx_content):
    result = []
    document = Document(io.BytesIO(docx_content))
    alignment_types = []
    isThereParagraph = False

    for para in document.paragraphs:
        if para.text != '':
            isThereParagraph = True
        if para.alignment is not None:
            alignment_types.append(para.alignment)
        text += para.text.strip() + '\n'
    unique_alignments = set(alignment_types)
    
    if len(unique_alignments) >= 2:
        result.append("The resume structure is multi-column. ATS may not be able to access the content in the multi-column structure.")
    
    if document.tables:
        if not isThereParagraph:
            result.append("The resume structure is table. ATS may not be able to access the content in the table structure.")
            return result
        result.append("The resume contains some table. ATS may not be able to access the content in the table.")

    return result